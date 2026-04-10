from pathlib import Path
from docx import Document

out = Path(r"E:\cli_planilhas\NVIDIA_NIM_Zed_Model_Guide.docx")

doc = Document()
doc.add_heading("NVIDIA NIM no Zed - Guia Pratico", 0)
doc.add_paragraph("Este guia cobre: campos da UI do Zed, quando usar cada tipo de modelo, modelos recomendados, e presets seguros para configuracao manual.")

doc.add_heading("1) Campos exatos da UI (como na imagem)", level=1)
for line in [
    "Model Name: ID exato do modelo (ex.: deepseek-ai/deepseek-v3.1-terminus).",
    "Max Completion Tokens: limite de geracao por resposta.",
    "Max Output Tokens: teto de saida em providers que respeitam esse campo.",
    "Max Tokens: contexto total (prompt + historico + output).",
    "Supports tools: marque para Agent usar ferramentas.",
    "Supports images: marque apenas para modelos vision.",
    "Supports parallel_tool_calls: geralmente OFF para compatibilidade.",
    "Supports prompt_cache_key: geralmente OFF.",
    "Supports /chat/completions: ON para NVIDIA NIM OpenAI-compatible."
]:
    doc.add_paragraph(line, style="List Bullet")

doc.add_heading("2) Presets recomendados para preencher rapido", level=1)
doc.add_paragraph("Preset padrao (texto/codigo):", style="List Bullet")
doc.add_paragraph("Max Completion Tokens = 32000 | Max Output Tokens = 32000 | Max Tokens = 200000", style="List Bullet 2")
doc.add_paragraph("Preset reasoning pesado:", style="List Bullet")
doc.add_paragraph("Mesmos valores; prefira modelos 'thinking' / 'reasoning' no nome.", style="List Bullet 2")
doc.add_paragraph("Preset vision:", style="List Bullet")
doc.add_paragraph("Mesmo token budget + Supports images = ON.", style="List Bullet 2")

doc.add_heading("3) Quando usar cada tipo de modelo", level=1)
rows = [
    ("Coding agent principal", "deepseek-ai/deepseek-v3.1-terminus, qwen/qwen3-coder-480b-a35b-instruct, mistralai/devstral-2-123b-instruct-2512", "Refactor, correcoes multi-arquivo, tool use."),
    ("Reasoning pesado", "moonshotai/kimi-k2-thinking, qwen/qwen3-next-80b-a3b-thinking, deepseek-ai/deepseek-r1-distill-qwen-32b", "Planejamento complexo, depuracao dificil, decisoes com trade-off."),
    ("Rapido/custo baixo", "openai/gpt-oss-20b, microsoft/phi-4-mini-instruct, meta/llama-3.1-8b-instruct", "Iteracao curta, classificacao, transformacoes simples."),
    ("Alta qualidade geral", "meta/llama-3.3-70b-instruct, mistralai/mistral-large-2-instruct, moonshotai/kimi-k2.5", "QA geral, escrita tecnica, assistente principal."),
    ("Vision multimodal", "meta/llama-3.2-90b-vision-instruct, meta/llama-3.2-11b-vision-instruct", "Explicar imagem/screenshot; OCR basico; validacao visual.")
]
for title, models, when in rows:
    doc.add_paragraph(f"{title}", style="List Number")
    doc.add_paragraph(f"Modelos: {models}", style="List Bullet 2")
    doc.add_paragraph(f"Use quando: {when}", style="List Bullet 2")

doc.add_heading("4) Modelos sugeridos (rank pratico para agent/coding)", level=1)
models = [
"deepseek-ai/deepseek-v3.1-terminus","moonshotai/kimi-k2-thinking","qwen/qwen3-coder-480b-a35b-instruct","mistralai/devstral-2-123b-instruct-2512","mistralai/mistral-small-4-119b-2603","meta/llama-3.3-70b-instruct","meta/llama-4-maverick-17b-128e-instruct","meta/llama-4-scout-17b-16e-instruct","qwen/qwen3-next-80b-a3b-thinking","qwen/qwen3-next-80b-a3b-instruct","qwen/qwen2.5-coder-32b-instruct","moonshotai/kimi-k2-instruct","moonshotai/kimi-k2.5","mistralai/mistral-large-2-instruct","mistralai/mistral-large","openai/gpt-oss-120b","openai/gpt-oss-20b","nvidia/llama-3.3-nemotron-super-49b-v1.5","nvidia/nemotron-3-super-120b-a12b","google/gemma-4-31b-it","meta/llama-3.2-90b-vision-instruct","meta/llama-3.2-11b-vision-instruct"
]
for m in models:
    doc.add_paragraph(m, style="List Bullet")

doc.add_heading("5) Checklist de validacao rapida", level=1)
for line in [
    "API URL do provider: https://integrate.api.nvidia.com/v1",
    "Model Name deve bater exatamente com /v1/models -> data[].id",
    "Se erro 401: problema de chave.",
    "Se erro model_not_found: ID errado.",
    "Se usar vision e falhar imagem: confirme Supports images = ON no modelo vision."
]:
    doc.add_paragraph(line, style="List Bullet")

doc.save(out)
print(out)
